import re
from docx import Document

try:
    doc = Document('template.docx')
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    placeholders = re.findall(r'\{\{(.*?)\}\}', text)
    print("--- PLACEHOLDERS FOUND ---")
    print(set(placeholders))
except Exception as e:
    print(f"Error: {e}")
