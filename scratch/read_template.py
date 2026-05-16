from docx import Document

try:
    doc = Document('template.docx')
    print("--- TEMPLATE CONTENT ---")
    for para in doc.paragraphs:
        print(para.text)
    print("--- TABLES ---")
    for table in doc.tables:
        print(f"Table with {len(table.rows)} rows and {len(table.columns)} columns")
except Exception as e:
    print(f"Error: {e}")
